import docx

docx_path = r'C:\Users\duyho\Downloads\TL\Mau Cuon Bao Cao Do An ĐBCLPM_Nhom10_EduShare_V7.docx'
out_path = r'C:\Users\duyho\Downloads\TL\Mau Cuon Bao Cao Do An ĐBCLPM_Nhom10_EduShare_V8.docx'

doc = docx.Document(docx_path)

ch4_content = [
    ("4.1. Môi trường và Công cụ phát triển", True, False),
    ("Dự án được xây dựng dựa trên các công nghệ và bộ công cụ hiện đại, tối ưu cho việc phát triển ứng dụng Web hiệu năng cao:", False, False),
    ("- Backend: Môi trường thực thi Node.js kết hợp framework Express.js để tạo ra các RESTful API phản hồi nhanh chóng.", False, False),
    ("- Cơ sở dữ liệu: Hệ quản trị cơ sở dữ liệu quan hệ MySQL. Sử dụng các truy vấn SQL thuần kết hợp Stored Procedures và Transaction để đảm bảo tính toàn vẹn dữ liệu.", False, False),
    ("- Frontend: Giao diện được code thuần bằng HTML5, CSS3 và Vanilla JavaScript. Có kết hợp thêm thư viện SweetAlert2 để tùy biến thông báo (Popup) và Chart.js để vẽ biểu đồ thống kê trực quan.", False, False),
    ("- Công cụ hỗ trợ: Visual Studio Code (Lập trình), Postman (Kiểm thử API độc lập), XAMPP (Máy chủ cơ sở dữ liệu cục bộ).", False, False),
    
    ("4.2. Kiến trúc và Tổ chức thư mục dự án", True, False),
    ("Dự án áp dụng triệt để mô hình Client - Server tách biệt, mã nguồn được chia thành hai phân hệ độc lập nhằm dễ dàng mở rộng và bảo trì:", False, False),
    ("- Phân hệ Client (Thư mục fe/):", True, False),
    ("  + Thư mục pages/: Chứa toàn bộ các giao diện HTML chia theo vai trò (admin, auth, document, group, guest, user).", False, False),
    ("  + Thư mục css/ và js/: Chứa các tệp định dạng và xử lý logic phía trình duyệt.", False, False),
    ("- Phân hệ Server (Thư mục be/):", True, False),
    ("  + Thư mục routes/: Định nghĩa các endpoint (API) của hệ thống.", False, False),
    ("  + Thư mục controllers/: Chứa logic xử lý nghiệp vụ cho từng endpoint.", False, False),
    ("  + Thư mục middlewares/: Xử lý các tiền điều kiện như Xác thực JWT, Kiểm tra quyền truy cập, và Rate Limiting chống Spam.", False, False),
    ("  + Thư mục config/: Chứa các cấu hình kết nối tới MySQL, Cloudinary, Nodemailer.", False, False),
    
    ("4.3. Tích hợp Dịch vụ Bên thứ 3 (Third-party APIs)", True, False),
    ("Để tăng cường trải nghiệm và bảo mật, hệ thống kết nối với các nền tảng dịch vụ nổi tiếng thông qua API:", False, False),
    ("- Cloudinary API: Toàn bộ ảnh đại diện (Avatar), ảnh bìa nhóm và tài liệu (PDF, Word) được upload thẳng lên hệ thống lưu trữ đám mây Cloudinary, giúp giảm tải bộ nhớ cho máy chủ vật lý.", False, False),
    ("- VirusTotal API: Tích hợp hệ thống quét mã độc tự động. Trước khi một tài liệu được lưu vào CSDL, file sẽ được băm SHA-256 và đẩy sang VirusTotal để kiểm tra tính an toàn.", False, False),
    ("- Google OAuth2 (Passport.js): Tích hợp tính năng 'Đăng nhập bằng Google', cho phép người dùng bỏ qua bước đăng ký rườm rà, hệ thống tự động trích xuất Email và tạo tài khoản ngay lập tức.", False, False),
    ("- Nodemailer: Dịch vụ SMTP tự động hóa việc gửi mã OTP khôi phục mật khẩu trực tiếp vào hòm thư điện tử của người dùng.", False, False),
    
    ("4.4. Các kỹ thuật Đảm bảo Chất lượng Mã nguồn (Code Quality & Security)", True, False),
    ("Đảm bảo chất lượng phần mềm không chỉ nằm ở khâu kiểm thử mà phải bắt đầu từ kỹ thuật viết code. Dự án EduShare áp dụng các chuẩn bảo mật sau:", False, False),
    ("- Xác thực Role-based Access Control (RBAC): Phân quyền nghiêm ngặt theo 3 cấp (SinhVien, GiaoVien, Admin). Toàn bộ Request phải đính kèm Access Token trong Header để Middleware giải mã và đối chiếu quyền.", False, False),
    ("- Bảo mật CSDL bằng Bcrypt & SHA-256: Mật khẩu người dùng được băm bằng thuật toán Bcrypt. Nội dung tài liệu được băm bằng SHA-256 để phát hiện và ngăn chặn việc tải lên trùng lặp (Đạo văn).", False, False),
    ("- Transaction Row-level Locking: Đây là kỹ thuật cốt lõi trong phân hệ Giao dịch Xu. Khi người dùng mua tài liệu, lệnh 'SELECT ... FOR UPDATE' được kích hoạt để khóa dòng dữ liệu tạm thời, triệt tiêu hoàn toàn lỗi âm tiền (Race Condition) trong trường hợp người dùng click chuột liên tục hoặc dùng tool tấn công.", False, False),
    ("- Chống Spam (Rate Limiting): Cấu hình express-rate-limit để khóa tạm thời các địa chỉ IP thực hiện quá số lượng truy vấn cho phép trong một thời gian ngắn, ngăn chặn nguy cơ DDoS làm sập hệ thống.", False, False),
]

insert_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "CHƯƠNG V. TRIỂN KHAI" in p.text.upper():
        insert_idx = i
        break

if insert_idx != -1:
    target_p = doc.paragraphs[insert_idx]
    
    for content, is_bold, is_italic in ch4_content:
        new_p = target_p.insert_paragraph_before(content)
        if is_bold:
            new_p.runs[0].bold = True
        if is_italic:
            new_p.runs[0].italic = True
    
    # Add an empty paragraph for spacing before CH5
    target_p.insert_paragraph_before("")

doc.save(out_path)
print("V8 generated successfully with Chapter IV content!")
