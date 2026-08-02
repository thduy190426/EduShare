import docx

docx_path = r'C:\Users\duyho\Downloads\TL\Mau Cuon Bao Cao Do An ĐBCLPM_Nhom10_EduShare_V8.docx'
out_path = r'C:\Users\duyho\Downloads\TL\Mau Cuon Bao Cao Do An ĐBCLPM_Nhom10_EduShare_V9.docx'

doc = docx.Document(docx_path)

missing_abbreviations = [
    ("RBAC", "Role-Based Access Control (Kiểm soát truy cập dựa trên vai trò)"),
    ("UML", "Unified Modeling Language (Ngôn ngữ mô hình hóa thống nhất)"),
    ("SQL", "Structured Query Language (Ngôn ngữ truy vấn mang tính cấu trúc)"),
    ("JSON", "JavaScript Object Notation (Định dạng trao đổi dữ liệu)"),
    ("HTML / CSS", "HyperText Markup Language / Cascading Style Sheets (Ngôn ngữ lập trình giao diện)"),
    ("THCS / THPT", "Trung học cơ sở / Trung học phổ thông"),
    ("SQA", "Software Quality Assurance (Đảm bảo chất lượng phần mềm)"),
    ("HTTP / HTTPS", "Hypertext Transfer Protocol Secure (Giao thức truyền tải siêu văn bản bảo mật)"),
    ("ACID", "Atomicity, Consistency, Isolation, Durability (Bốn thuộc tính của một giao dịch dữ liệu an toàn)"),
    ("SMTP", "Simple Mail Transfer Protocol (Giao thức truyền tải thư tín đơn giản)"),
    ("IP", "Internet Protocol (Giao thức mạng Internet)")
]

insert_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("OAuth2:"):
        insert_idx = i
        break

if insert_idx != -1:
    target_p = doc.paragraphs[insert_idx]
    
    for abbr, desc in missing_abbreviations:
        new_p = target_p.insert_paragraph_before("")
        new_p.style = doc.styles['List Bullet'] if 'List Bullet' in doc.styles else doc.styles['Normal']
        
        run_abbr = new_p.add_run(f"{abbr}")
        run_abbr.bold = True
        
        new_p.add_run(f": {desc}")

doc.save(out_path)
print("V9 generated successfully with additional Abbreviations!")
